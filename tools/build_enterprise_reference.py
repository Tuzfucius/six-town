"""从调研底稿生成供页面扩充使用的企业基础信息参考册。"""

from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"E:\桌面\社会实践\后哨工作\企业.docx")
OUTPUT = ROOT / "docs" / "六镇相关企业基础信息参考册.docx"

TOWNS = {
    "2": "梦想小镇",
    "3": "人工智能小镇/未来科技城",
    "4": "中国（杭州）算力小镇",
    "5": "德清地理信息小镇",
    "6": "越城集成电路小镇",
    "7": "嘉善归谷智造小镇",
}
SECTION_RE = re.compile(r"^([2-7])\.(\d+)\s+(.+)$")
SUBSECTION_RE = re.compile(r"^[2-7]\.\d+\.(\d+)\s+(.+)$")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill)
    tc_pr.append(shade)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths=(2700, 6660)) -> None:
    table.autofit = False
    table.style = "Table Grid"
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for col, width in zip(grid.gridCol_lst, widths):
        col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            tc_pr = cell._tc.get_or_add_tcPr()
            margins = OxmlElement("w:tcMar")
            for side in ("top", "start", "bottom", "end"):
                margin = OxmlElement(f"w:{side}")
                margin.set(qn("w:w"), "120" if side in ("start", "end") else "80")
                margin.set(qn("w:type"), "dxa")
                margins.append(margin)
            tc_pr.append(margins)


def add_text(doc: Document, text: str, style: str = "Normal", bold_label: str | None = None) -> None:
    paragraph = doc.add_paragraph(style=style)
    if bold_label and text.startswith(bold_label):
        run = paragraph.add_run(bold_label)
        run.bold = True
        paragraph.add_run(text[len(bold_label):])
    else:
        paragraph.add_run(text)


def source_records(source: Document):
    tables = [table for table in source.tables if len(table.rows) == 6 and len(table.columns) == 2]
    table_records = []
    for table in tables:
        row_data = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in table.rows}
        table_records.append(row_data)

    records = []
    current = None
    current_subsection = None
    for paragraph in source.paragraphs:
        text = paragraph.text.strip()
        section_match = SECTION_RE.match(text)
        if section_match:
            chapter, number, title = section_match.groups()
            if int(number) >= 2:
                current = {
                    "chapter": chapter,
                    "title": title,
                    "fields": defaultdict(list),
                }
                records.append(current)
            else:
                current = None
            current_subsection = None
            continue
        if current is None:
            continue
        subsection_match = SUBSECTION_RE.match(text)
        if subsection_match:
            current_subsection = subsection_match.group(2)
            continue
        if text and current_subsection:
            current["fields"][current_subsection].append(text)

    records = [
        record
        for record in records
        if "企业基本信息" not in record["title"]
        and "产业生态" not in record["title"]
        and "科研机构与公共平台" not in record["title"]
    ]
    if len(records) != len(table_records):
        raise RuntimeError(f"企业章节数 {len(records)} 与信息表数 {len(table_records)} 不一致")
    for record, data in zip(records, table_records):
        record["meta"] = data
    return records


def make_document(records) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in (("Heading 1", 16, "2E74B5"), ("Heading 2", 13, "2E74B5"), ("Heading 3", 11.5, "1F4D78")):
        style = doc.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 3" else 8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.add_run("六镇相关企业基础信息参考册 | 页面内容扩充底稿").font.size = Pt(8.5)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("浙江工业大学健行学院实践团 | 基于 2026 年 7 月第一阶段材料整理").font.size = Pt(8.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(32)
    title.paragraph_format.space_after = Pt(10)
    run = title.add_run("六镇相关企业基础信息参考册")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("用于数字图谱平台页面内容扩充、企业卡片配置与后续实地核验").font.size = Pt(12)
    add_text(doc, "资料口径：本册由《杭湖嘉绍都市圈六镇企业与产业材料汇总（第一阶段正式稿 V1.0）》整理而成，覆盖 29 个企业或项目主体。所有“待核验”“部分核验”内容仅可作为页面候选信息，不得作为确定性事实对外发布。")
    doc.add_paragraph().add_run("使用建议：页面数据至少保留企业名称、企业 ID、所属小镇、产业标签、产业链位置、地址性质、核验状态、代表产品、页面摘要和来源字段；产品参数、地址坐标和经营状态应在发布前复核。")

    summary = doc.add_table(rows=1, cols=2)
    configure_table(summary)
    summary.cell(0, 0).text = "小镇"
    summary.cell(0, 1).text = "纳入企业/项目主体数量"
    for cell in summary.rows[0].cells:
        set_cell_shading(cell, "E8EEF5")
        for run in cell.paragraphs[0].runs:
            run.bold = True
    counts = defaultdict(int)
    for record in records:
        counts[TOWNS[record["chapter"]]] += 1
    for town in TOWNS.values():
        cells = summary.add_row().cells
        cells[0].text = town
        cells[1].text = str(counts[town])

    fields = [
        ("2", "产业与产业链位置"),
        ("3", "地址与小镇关联"),
        ("4", "代表产品、技术或成果"),
        ("5", "企业特色与竞争优势"),
        ("6", "新质生产力体现"),
        ("7", "对小镇产业生态的作用"),
        ("8", "数据来源与待核验事项"),
    ]
    current_town = None
    for record in records:
        town = TOWNS[record["chapter"]]
        if town != current_town:
            if current_town is not None:
                doc.add_page_break()
            doc.add_heading(town, level=1)
            add_text(doc, "以下条目按原始调研材料的企业顺序整理。页面展示时应保留核验状态，并将地址、企业主体与小镇归属作为独立字段维护。")
            current_town = town

        doc.add_heading(record["title"], level=2)
        meta_table = doc.add_table(rows=0, cols=2)
        configure_table(meta_table)
        for label in ("企业ID", "企业类型", "小队联系标识", "企业名称", "地址", "官网", "核验状态"):
            value = record["meta"].get(label, "待补充")
            row = meta_table.add_row().cells
            row[0].text = label
            row[1].text = value or "待补充"
            set_cell_shading(row[0], "F2F4F7")
            for run in row[0].paragraphs[0].runs:
                run.bold = True

        for key, label in fields:
            content = " ".join(record["fields"].get(key, [])) or "待补充"
            doc.add_heading(label, level=3)
            add_text(doc, content)

        note = doc.add_paragraph()
        note.paragraph_format.space_before = Pt(4)
        note_run = note.add_run("页面配置提示：")
        note_run.bold = True
        note_run.font.color.rgb = RGBColor(122, 90, 0)
        note.add_run("优先展示已核验的企业全称、产业标签、代表产品与来源；地址、量产状态、性能参数和小镇核心区归属在核验完成前应展示为“待核验”或不在前台展示。")

    doc.add_page_break()
    doc.add_heading("数据维护规范", level=1)
    for item in (
        "企业地址应拆分为注册地址、实际办公地、研发地、生产基地和历史孵化地；地图点位应记录地址性质、经纬度、精度和核验来源。",
        "产业信息应拆分为一级标签、二级标签和产业链位置。企业可有多个标签，但页面主标签须与统计口径一致。",
        "成果状态统一使用“已发布”“已获批/认证”“已量产”“示范应用”“在研”“规划建设”或“待核验”，避免以宣传性表述替代事实状态。",
        "每次更新应记录来源链接、获取日期、核验人、证据摘要和版本号；涉及医疗器械、芯片产能与技术参数的内容需回到原始材料复核。",
    ):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.add_run(item)
    return doc


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(f"未找到源文档：{SOURCE}")
    OUTPUT.parent.mkdir(exist_ok=True)
    records = source_records(Document(SOURCE))
    make_document(records).save(OUTPUT)
    print(f"已生成：{OUTPUT}")
    print(f"企业条目：{len(records)}")


if __name__ == "__main__":
    main()
